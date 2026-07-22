#!/usr/bin/env python3
"""Generate Chinese Word reports for the KDebug repair benchmark."""

import argparse
import csv
import statistics
from collections import defaultdict
from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


MODELS = ["gpt-5.5", "glm-4.7", "qwen3.6-35b"]
GROUP_LABEL = {
    "with_kdebug": "使用 kdebug",
    "without_kdebug": "不使用 kdebug",
}
STATUS_LABEL = {
    "PASS": "通过",
    "TIMEOUT": "未修通（1 小时超时）",
    "RULE_VIOLATION": "规则违规/无效",
    "TOOL_EVIDENCE_MISSING": "工具证据缺失",
    "TOOL_EVIDENCE_INVALID": "工具证据无效",
    "INFRA_ERROR": "基础设施异常",
    "FAIL": "基础设施异常",
    "RETRY_LATER": "限流/接口异常，待重试",
    "missing": "缺失",
}
BUG_DOMAIN_LABEL = {
    "rtl": "RTL 注错",
    "env": "环境注错",
    "mixed": "RTL + 环境混合注错",
}
LAYER_LABEL = {
    "generated_wrapper": "Generated RTL wrapper",
    "real_rtl_ut": "真实 RTL UT",
    "real_rtl_it": "真实 RTL IT",
    "fullchip": "整芯片 XiangShan",
}

CASE_DETAILS = {
    "case_001": "AXI ready/valid 写通道握手在 backpressure 下被条件性破坏，表现为 AW/W/B 顺序或响应配对错误。",
    "case_002": "AXI error/backpressure 场景下的 burst、last beat 或 response status 传播异常，scoreboard 会看到错误响应被吞掉或变成 OK。",
    "case_003": "Memory 子系统地址、byte mask 或 data beat 被破坏，读回数据在特定 lane 或 beat 上不一致。",
    "case_004": "MMU 或地址权限/翻译边界处理异常，合法与非法访问交错时状态返回不符合设计约束。",
    "case_005": "Peripheral 子系统 error response 在 ready 延迟下发生丢失、重复或乱序。",
    "case_006": "Late pipeline bug，运行数千到数万周期后才由 difftest 或 workload 暴露，早期寄存器状态可能完全正常。",
    "case_007": "Control bug，valid/ready/rfWen/flush 等控制信号被条件性破坏，根因通常早于最终 mismatch 出现的位置。",
    "case_008": "LSU/cache bug，地址、mask、data beat 或 response 顺序异常，最终失败点可能距离 first bad event 很远。",
    "case_009": "Branch/redirect bug，PC 发生错误但寄存器暂时不 mismatch，需要追踪 fetch、redirect 与 commit PC 的时间关系。",
    "case_010": "真实 RTL Cache/MMU/refill/permission 压测，要求在真实 RTL 上完成闭环修复。",
    "case_011": "环境 bug，run option 丢失 difftest/REF_SO 等关键参数，仿真可能执行但判定证据无效。",
    "case_012": "环境 bug，run dispatch 把 case、seed 或 workload 指向错误压测，失败现象与 metadata 不一致。",
    "case_013": "环境 bug，build/run 使用 stale simv、错误 out 路径或执行权限异常，RTL 修改未进入最终仿真。",
    "case_014": "混合 bug，memory mask/data beat RTL 错误叠加 run/judge 日志路径或 pass marker 环境错误。",
    "case_015": "混合 bug，branch/redirect RTL 错误叠加错误 REF_SO 或 run option，环境噪声会污染失败表现。",
    "case_016": "混合 bug，LSU/cache RTL 错误叠加过短 timeout 或错误 workload 参数，需要同时修 RTL 和环境。",
}


def read_rows(path):
    with open(path, newline="", encoding="utf-8") as f:
        return list(csv.DictReader(f))


def as_float(value, default=0.0):
    try:
        return float(value)
    except (TypeError, ValueError):
        return default


def as_int(value):
    return int(as_float(value))


def median(values):
    nums = [as_float(v) for v in values if str(v) != ""]
    return statistics.median(nums) if nums else 0.0


def status_text(value):
    return STATUS_LABEL.get(value, value or "")


def setup_document(doc):
    section = doc.sections[0]
    section.top_margin = Inches(0.72)
    section.bottom_margin = Inches(0.72)
    section.left_margin = Inches(0.76)
    section.right_margin = Inches(0.76)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Microsoft YaHei"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    normal.font.size = Pt(9.5)
    normal.paragraph_format.space_after = Pt(4)
    normal.paragraph_format.line_spacing = 1.08

    for name, size in [("Heading 1", 14.5), ("Heading 2", 12), ("Heading 3", 10.8)]:
        style = styles[name]
        style.font.name = "Microsoft YaHei"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor(31, 78, 121)
        style.paragraph_format.space_before = Pt(8)
        style.paragraph_format.space_after = Pt(4)


def set_run_font(run, size=None, bold=None, color=None):
    name = "Microsoft YaHei"
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:ascii"), name)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), name)
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if color is not None:
        run.font.color.rgb = RGBColor(*color)


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=110, bottom=80, end=110):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def add_title(doc, title, subtitle=None):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(title)
    set_run_font(r, size=17, bold=True, color=(11, 37, 69))
    if subtitle:
        p2 = doc.add_paragraph()
        p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r2 = p2.add_run(subtitle)
        set_run_font(r2, size=9.5, color=(85, 85, 85))


def add_para(doc, text, bold=False, color=None):
    p = doc.add_paragraph()
    r = p.add_run(str(text))
    set_run_font(r, bold=bold, color=color)
    return p


def add_table(doc, headers, rows):
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for i, header in enumerate(headers):
        hdr[i].text = ""
        p = hdr[i].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(str(header))
        set_run_font(r, size=8, bold=True)
        hdr[i].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        set_cell_shading(hdr[i], "F2F4F7")
        set_cell_margins(hdr[i])
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            text = str(value)
            cells[i].text = ""
            p = cells[i].paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT if len(text) > 14 else WD_ALIGN_PARAGRAPH.CENTER
            r = p.add_run(text)
            set_run_font(r, size=7.6)
            cells[i].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cells[i])
    doc.add_paragraph()
    return table


def stats(rows):
    total = len(rows)
    passed = sum(1 for r in rows if r.get("final_status") == "PASS")
    timeout = sum(1 for r in rows if r.get("final_status") == "TIMEOUT")
    missing = sum(1 for r in rows if r.get("final_status") == "TOOL_EVIDENCE_MISSING")
    invalid = sum(1 for r in rows if r.get("final_status") == "TOOL_EVIDENCE_INVALID")
    rule = sum(1 for r in rows if r.get("final_status") == "RULE_VIOLATION")
    retry = sum(1 for r in rows if r.get("final_status") == "RETRY_LATER")
    infra = total - passed - timeout - missing - invalid - rule - retry
    valid_total = passed + timeout
    success_rate = f"{(passed / valid_total * 100):.1f}%" if valid_total else "0.0%"
    pass_median = (
        f"{median(r.get('elapsed_sec') for r in rows if r.get('final_status') == 'PASS'):.1f}"
        if passed else ""
    )
    return [
        total, passed, timeout, missing, invalid, rule, retry, infra, success_rate,
        pass_median, str(sum(as_int(r.get("token_total")) for r in rows)),
    ]


def grouped_summary_rows(rows, fields):
    groups = defaultdict(list)
    for row in rows:
        key = tuple(row.get(field, "") for field in fields)
        groups[key].append(row)
    return [list(key) + stats(groups[key]) for key in sorted(groups)]


def screenshot_candidates(screenshot_root, row):
    case_id = row.get("case_id", "")
    model = row.get("model_id", "")
    group = row.get("group", "")
    return [
        screenshot_root / f"{case_id}_{model}_{group}.png",
        screenshot_root / f"{model}_{group}_{case_id}.png",
        screenshot_root / case_id / f"{model}_{group}.png",
        screenshot_root / model / group / f"{case_id}.png",
    ]


def add_screenshot(doc, screenshot_root, row):
    found = next((p for p in screenshot_candidates(screenshot_root, row) if p.exists()), None)
    add_para(doc, "终端截图证据：", bold=True)
    if not found:
        add_para(
            doc,
            "截图缺失：该 trial 不满足最终交付证据要求。",
            bold=True,
            color=(192, 0, 0),
        )
        return
    try:
        doc.add_picture(str(found), width=Inches(6.65))
        cap = doc.add_paragraph()
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = cap.add_run(str(found))
        set_run_font(r, size=7, color=(100, 100, 100))
    except Exception as exc:
        add_para(doc, f"截图插入失败：{found}，原因：{exc}", bold=True, color=(192, 0, 0))


def result_brief(row):
    if not row:
        return "缺失"
    return (
        f"{status_text(row.get('final_status'))}，耗时 {row.get('elapsed_sec', '')}s，"
        f"迭代 {row.get('iterations', '')}，Token {row.get('token_total', '')}，"
        f"修复类型 {row.get('repair_class', '')}"
    )


def bug_domain_text(row):
    return BUG_DOMAIN_LABEL.get(row.get("bug_domain", ""), row.get("bug_domain", ""))


def layer_text(row):
    layer = row.get("benchmark_layer", "") or row.get("layer", "")
    return LAYER_LABEL.get(layer, layer)


def trial_detail_rows(row):
    return [
        ["Case", row.get("case_id", "")],
        ["注错域", bug_domain_text(row)],
        ["Benchmark 层级", layer_text(row)],
        ["子系统", row.get("subsystem", "")],
        ["故障类别", row.get("bug_class", "")],
        ["环境故障类别", row.get("env_fault_class", "")],
        ["最终状态", status_text(row.get("final_status", ""))],
        ["失败/异常分类", row.get("failure_class", "")],
        ["总耗时(s)", row.get("elapsed_sec", "")],
        [
            "定位/修改/构建/运行/判定(s)",
            f"{row.get('locate_sec','')}/{row.get('edit_sec','')}/"
            f"{row.get('build_sec','')}/{row.get('run_sec','')}/{row.get('judge_sec','')}",
        ],
        ["迭代次数", row.get("iterations", "")],
        ["Token", row.get("token_total", "")],
        ["RTL 修改", row.get("modified_rtl_files", "")],
        ["环境修改", row.get("modified_env_files", "")],
        [
            "工具证据",
            f"required={row.get('tool_evidence_required','')}, "
            f"present={row.get('tool_evidence_present','')}, "
            f"valid={row.get('tool_evidence_valid','')}, "
            f"collection={row.get('tool_evidence_collection_id','')}, "
            f"files={row.get('tool_evidence_files','')}",
        ],
        ["公开设计资料", row.get("public_design_refs", "")],
        ["Pass marker", row.get("pass_marker", "")],
        ["规则/失败说明", row.get("rule_violation", "")],
    ]


def add_protocol_section(doc):
    doc.add_heading("一、测试口径", level=1)
    add_para(
        doc,
        "每个 case、每个模型、每个工具组最多 3600 秒。只有模型修改允许范围内的 RTL 或环境文件，"
        "重新 build/run 原始 fail workload，并由受保护 judge 判定通过，才记为 PASS。",
    )
    add_para(
        doc,
        "repair loop 中出现 no patch、build fail、run fail、judge fail 时，runner 会把日志继续反馈给模型重修；"
        "这些只是中间状态，不是最终失败。普通模型失败严格等于 1 小时预算耗尽仍未 PASS，即 TIMEOUT。",
    )
    add_para(
        doc,
        "API 限流、接口无响应等单次尝试先单列为 RETRY_LATER；但同一模型、工具组、case 多次 RETRY_LATER "
        "累计耗时超过 1 小时后，也按预算耗尽记为 TIMEOUT。runner 崩溃、脚本兼容性、截图/Word 生成异常等"
        "仍属于基础设施问题，单列为 INFRA_ERROR。",
    )
    add_para(
        doc,
        "with_kdebug 组必须通过独立 KDebug 采集并持有可复核 manifest；缺失标记 TOOL_EVIDENCE_MISSING，"
        "manifest、哈希、输入或响应校验失败标记 TOOL_EVIDENCE_INVALID，二者都不能用于证明工具效果。",
    )


def summary_headers(prefix):
    return prefix + [
        "Trial",
        "PASS",
        "未修通=TIMEOUT",
        "证据缺失",
        "证据无效",
        "规则违规",
        "限流待重试",
        "基础设施异常",
        "有效成功率",
        "PASS中位耗时",
        "Token总量",
    ]


def generate_model_report(rows, model, out_path, screenshot_root):
    doc = Document()
    setup_document(doc)
    add_title(doc, f"{model} XiangShan Debug Benchmark 报告", "with kdebug / without kdebug repair loop 对比")
    add_protocol_section(doc)

    model_rows = [r for r in rows if r.get("model_id") == model]
    doc.add_heading("二、模型汇总", level=1)
    add_table(doc, summary_headers(["工具组"]), grouped_summary_rows(model_rows, ["group"]))

    doc.add_heading("三、按注错域拆分", level=1)
    add_table(doc, summary_headers(["注错域", "工具组"]), grouped_summary_rows(model_rows, ["bug_domain", "group"]))

    doc.add_heading("四、逐 Case 结果与截图", level=1)
    for row in sorted(model_rows, key=lambda r: (r.get("case_id", ""), r.get("group", ""))):
        case_id = row.get("case_id", "")
        group_label = GROUP_LABEL.get(row.get("group"), row.get("group"))
        doc.add_heading(f"{case_id} - {group_label}", level=2)
        add_para(doc, CASE_DETAILS.get(case_id, "该 case 的公开注错说明未在脚本中登记；以 case_meta 与日志为准。"))
        add_table(doc, ["字段", "值"], trial_detail_rows(row))
        add_screenshot(doc, screenshot_root, row)

    doc.save(out_path)


def missing_screenshots(rows, screenshot_root):
    missing = []
    for row in rows:
        if not any(p.exists() for p in screenshot_candidates(screenshot_root, row)):
            missing.append(f"{row.get('case_id')} / {row.get('model_id')} / {row.get('group')}")
    return missing


def generate_summary_report(rows, out_path, screenshot_root):
    doc = Document()
    setup_document(doc)
    add_title(doc, "三模型 XiangShan Debug Benchmark 汇总报告", "gpt-5.5 / glm-4.7 / qwen3.6-35b")
    add_protocol_section(doc)

    doc.add_heading("二、总体结果", level=1)
    add_table(doc, summary_headers(["模型", "工具组"]), grouped_summary_rows(rows, ["model_id", "group"]))

    doc.add_heading("三、按注错域结果", level=1)
    add_table(doc, summary_headers(["注错域", "模型", "工具组"]), grouped_summary_rows(rows, ["bug_domain", "model_id", "group"]))

    doc.add_heading("四、按 Benchmark 层级结果", level=1)
    add_table(doc, summary_headers(["Benchmark 层级", "模型", "工具组"]), grouped_summary_rows(rows, ["benchmark_layer", "model_id", "group"]))

    doc.add_heading("五、逐 Case 对比", level=1)
    index = {(r.get("case_id"), r.get("model_id"), r.get("group")): r for r in rows}
    compare_rows = []
    for case_id in sorted({r.get("case_id", "") for r in rows if r.get("case_id")}):
        first = next(r for r in rows if r.get("case_id") == case_id)
        for model in MODELS:
            compare_rows.append([
                case_id,
                bug_domain_text(first),
                layer_text(first),
                first.get("subsystem", ""),
                model,
                result_brief(index.get((case_id, model, "with_kdebug"))),
                result_brief(index.get((case_id, model, "without_kdebug"))),
            ])
    add_table(doc, ["Case", "注错域", "Benchmark 层级", "子系统", "模型", "使用 kdebug", "不使用 kdebug"], compare_rows)

    doc.add_heading("六、截图完整性", level=1)
    missing = missing_screenshots(rows, screenshot_root)
    if missing:
        add_para(doc, "以下 trial 缺少终端截图，报告不满足最终交付要求：", bold=True, color=(192, 0, 0))
        for item in missing:
            add_para(doc, item)
    else:
        add_para(doc, "所有 trial 均找到终端截图。")

    doc.save(out_path)


def main():
    parser = argparse.ArgumentParser()
    parser.add_argument("--results", required=True, type=Path)
    parser.add_argument("--screenshots", required=True, type=Path)
    parser.add_argument("--out-dir", required=True, type=Path)
    parser.add_argument("--require-screenshots", action="store_true")
    args = parser.parse_args()

    rows = read_rows(args.results)
    args.out_dir.mkdir(parents=True, exist_ok=True)
    for model in MODELS:
        generate_model_report(rows, model, args.out_dir / f"{model}_benchmark_report.docx", args.screenshots)
    generate_summary_report(rows, args.out_dir / "three_model_summary_report.docx", args.screenshots)
    if args.require_screenshots:
        missing = missing_screenshots(rows, args.screenshots)
        if missing:
            raise SystemExit("missing screenshots: " + "; ".join(missing))


if __name__ == "__main__":
    main()
