#!/usr/bin/env python3
"""Generate a Chinese DOCX describing v2 private fault-injection intent.

This document is for benchmark operators.  It describes what to construct, not
the exact private injected line numbers.  Do not copy it into model-visible
case directories.
"""

import argparse
import csv
from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


DETAILS = {
    "ready_valid_aw_w_b_order": [
        "AXI 写地址、写数据、写响应配对在 backpressure 和多 outstanding 下被条件性破坏。",
        "fail log 可暴露 scoreboard mismatch，但不能暴露具体触发 bit、队列名或注错表达式。",
        "with_kdebug 证据应包含 AW/W/B 时间窗、ID/owner 对应关系和 first bad response。",
    ],
    "burst_len_or_last": [
        "AXI burst 长度归一化或 last beat 生成在边界组合下异常。",
        "失败应表现为 beat 数、last 标记或 response status 与 reference 不一致。",
        "with_kdebug 证据应包含 AWLEN、beat index、WLAST/RLAST 与 first mismatch。",
    ],
    "address_or_mask": [
        "memory path 地址低位、byte mask 或 data beat index 被条件性破坏。",
        "失败应稳定落在局部 byte lane、局部 beat 或非对齐访问上。",
        "with_kdebug 证据应串起 write request、mask/data beat 和 readback mismatch。",
    ],
    "mmio_resp_or_decode": [
        "MMIO decode 或 response status 在合法/非法访问交错时异常。",
        "失败应为 target/status/data mismatch，而不是全局 timeout。",
        "with_kdebug 证据应包含地址、decode target、response status 与 scoreboard 对比。",
    ],
    "error_backpressure": [
        "peripheral error response 在 ready 延迟窗口内发生 payload 不稳定、重复、丢失或错序。",
        "失败应能稳定复现同一类 response order/status 问题。",
        "with_kdebug 证据应覆盖 resp_valid && !resp_ready 窗口中的 payload hold。",
    ],
    "late_pipeline_bug": [
        "full-chip late pipeline bug 必须在数千到数万周期后失败，不能启动即失败。",
        "可破坏 result、bypass、writeback 或提交相关路径中的局部条件。",
        "with_kdebug 证据应从最终 difftest mismatch 回溯到 first divergent commit/writeback event。",
    ],
    "valid_ready_rfWen_flush": [
        "control bug 只破坏 valid、ready、rfWen 或 flush 等门控/控制信号，数据本身可以暂时正确。",
        "失败应体现控制时序错位、提交错误或写回被吞。",
        "with_kdebug 证据应证明 first bad event 在控制链路，而不是最终寄存器 mismatch 附近。",
    ],
    "address_mask_data_resp": [
        "LSU/cache bug 破坏地址、mask、data beat 或 response 顺序。",
        "根因和最终 load/store mismatch 可相隔很远。",
        "with_kdebug 证据应包含 load/store request、DCache/MSHR/refill source、beat 与最终 writeback data。",
    ],
    "pc_redirect_without_early_rf_mismatch": [
        "branch/redirect bug 破坏 PC 或 redirect target，但寄存器不一定立刻 mismatch。",
        "失败可先表现为 PC mismatch、取指路径异常或稍后的 architectural mismatch。",
        "with_kdebug 证据应覆盖 redirect producer、frontend target 和 commit PC。",
    ],
    "cache_mmu_interaction": [
        "Cache/MMU bug 破坏 VA/PA、permission、TLB/PTW response 或 refill 对应关系。",
        "失败应依赖 translation 与 cache/refill 压力，而不是简单 page fault。",
        "with_kdebug 证据应覆盖 VA/PA、permission、DCache request、MSHR/refill 链路。",
    ],
    "dropped_diff_plusarg": [
        "环境 bug：run 脚本条件性丢失 +diff/REF_SO，导致仿真执行但 difftest 无效。",
        "修复应修改 scripts/run.sh、env 或 config，不能靠改 RTL 让 judge 误判。",
        "with_kdebug 证据应包含 run command audit 和 difftest enabled 检查。",
    ],
    "wrong_case_or_seed": [
        "环境 bug：run dispatch 把目标 UVM case 或 seed 指向错误压测。",
        "模型必须识别 metadata、fail log 和 run script 的不一致。",
        "with_kdebug 证据应包含 case dispatch audit 和 scoreboard case map。",
    ],
    "stale_or_unwritable_simv": [
        "环境 bug：build/run 使用 stale simv、错误 out 路径或不可执行权限。",
        "模型必须修复构建产物路径、权限或 cache invalidation，使修改进入最终仿真。",
        "with_kdebug 证据应包含 build fingerprint、simv timestamp 和运行路径审计。",
    ],
    "rtl_memory_mask_plus_wrong_judge": [
        "混合 bug：memory mask/data beat RTL 错误，同时 run/judge 日志路径或 marker 配置错误。",
        "成功必须同时修改 RTL 与环境文件。",
        "with_kdebug 证据应分别指向 data mismatch 和环境路径/marker 异常。",
    ],
    "rtl_redirect_plus_bad_ref": [
        "混合 bug：branch/redirect RTL 错误叠加错误 REF_SO 或 run option。",
        "环境噪声应让单纯静态 RTL 定位更难，但不能让 case 变成不可判定。",
        "with_kdebug 证据应包含 redirect trace 和 REF_SO/run option audit。",
    ],
    "rtl_lsu_plus_timeout_config": [
        "混合 bug：LSU/cache RTL 错误叠加过短 timeout 或错误 workload 参数。",
        "模型必须修复 RTL 根因，并修正环境配置让 workload 能跑到有效 pass/fail。",
        "with_kdebug 证据应包含 LSU/cache first bad event 和 timeout/workload 配置审计。",
    ],
}


def set_run_font(run, name="Microsoft YaHei", size=None, bold=None, color=None):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:ascii"), name)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), name)
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if color is not None:
        run.font.color.rgb = RGBColor(*color)


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def setup(doc):
    section = doc.sections[0]
    section.top_margin = Inches(0.8)
    section.bottom_margin = Inches(0.8)
    section.left_margin = Inches(0.85)
    section.right_margin = Inches(0.85)
    normal = doc.styles["Normal"]
    normal.font.name = "Microsoft YaHei"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    normal.font.size = Pt(10)
    normal.paragraph_format.space_after = Pt(5)
    for name, size in [("Heading 1", 15), ("Heading 2", 12)]:
        style = doc.styles[name]
        style.font.name = "Microsoft YaHei"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor(46, 116, 181)


def para(doc, text, bold=False, color=None):
    p = doc.add_paragraph()
    r = p.add_run(text)
    set_run_font(r, bold=bold, color=color)


def bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    r = p.add_run(text)
    set_run_font(r)


def table(doc, headers, rows):
    t = doc.add_table(rows=1, cols=len(headers))
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    t.style = "Table Grid"
    for i, header in enumerate(headers):
        cell = t.rows[0].cells[i]
        cell.text = ""
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(header)
        set_run_font(r, size=8.5, bold=True)
        set_cell_shading(cell, "F2F4F7")
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    for row in rows:
        cells = t.add_row().cells
        for i, value in enumerate(row):
            cells[i].text = ""
            p = cells[i].paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            r = p.add_run(str(value))
            set_run_font(r, size=8)
            cells[i].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    doc.add_paragraph()


def read_case_matrix(path):
    if not path.exists():
        return []
    with path.open(newline="", encoding="utf-8") as f:
        return list(csv.DictReader(f))


def main():
    parser = argparse.ArgumentParser()
    parser.add_argument("--out", required=True, type=Path)
    parser.add_argument("--case-matrix", type=Path, default=Path(__file__).resolve().parents[1] / "case_matrix.csv")
    args = parser.parse_args()
    args.out.parent.mkdir(parents=True, exist_ok=True)

    rows = read_case_matrix(args.case_matrix)
    doc = Document()
    setup(doc)
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = title.add_run("KDebug XiangShan Benchmark v2 注错内容说明")
    set_run_font(r, size=18, bold=True, color=(11, 37, 69))
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = subtitle.add_run("私有构造资料：不得放入模型可见目录")
    set_run_font(r, size=10, color=(150, 0, 0))

    doc.add_heading("一、使用边界", level=1)
    para(doc, "本文件用于 benchmark 构造者规划私有注错和 answer_key_private.json。它不能复制到 case 目录、repair 目录、模型 prompt 或公开报告中。", bold=True, color=(150, 0, 0))
    bullet(doc, "RTL 中不能保留 bug/fault/inject/answer/kverif 等泄露命名或注错注释。")
    bullet(doc, "公开 case 只能使用 case_001 这类匿名目录，公开 metadata 不能包含真实注错点。")
    bullet(doc, "with_kdebug 证据可以给动态观测结果，但不能直接给 private answer key。")

    doc.add_heading("二、逐 Case 注错设计", level=1)
    for row in rows:
        case_id = row.get("case_id", "")
        bug_class = row.get("bug_class", "")
        doc.add_heading(f"{case_id}：{row.get('bug_domain', '')} / {row.get('subsystem', '')}", level=2)
        detail = DETAILS.get(bug_class, ["该 case 应遵循 case_matrix.csv 中的中文 notes 和 v2 repair scope。"])
        table(
            doc,
            ["字段", "内容"],
            [
                ["Benchmark 层", row.get("benchmark_layer", "")],
                ["目标 flow", row.get("target_flow", "")],
                ["Fail workload/case", row.get("fail_workload_or_case", "")],
                ["修复范围", row.get("repair_scope", "")],
                ["环境故障类别", row.get("env_fault_class", "")],
                ["工具证据要求", row.get("expected_evidence", "")],
                ["中文注错内容", "\n".join(detail)],
                ["公开边界", row.get("notes", "")],
            ],
        )

    doc.add_heading("三、验收检查", level=1)
    bullet(doc, "干净版本必须先 pass，注错版本必须稳定 fail。")
    bullet(doc, "env case 的正确修复必须体现在 scripts/env/config/filelists 等环境文件中。")
    bullet(doc, "mixed case 的正确修复必须同时体现在 RTL 和环境文件中。")
    bullet(doc, "with_kdebug 组必须先完成逐 case 独立 KDebug 采集并通过 manifest 校验；缺失标记 TOOL_EVIDENCE_MISSING，来源、输入、哈希、调用或响应无效标记 TOOL_EVIDENCE_INVALID。")

    doc.save(args.out)


if __name__ == "__main__":
    main()
