#!/usr/bin/env python3
"""Generate a Chinese DOCX design brief for benchmark v2."""

import argparse
import csv
from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


def set_run_font(run, name="Microsoft YaHei", size=None, bold=None, color=None):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:ascii"), name)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), name)
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if color is not None:
        run.font.color.rgb = RGBColor(*color)


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def setup(doc):
    section = doc.sections[0]
    section.top_margin = Inches(0.8)
    section.bottom_margin = Inches(0.8)
    section.left_margin = Inches(0.85)
    section.right_margin = Inches(0.85)
    normal = doc.styles["Normal"]
    normal.font.name = "Microsoft YaHei"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    normal.font.size = Pt(10)
    normal.paragraph_format.space_after = Pt(5)
    for name, size in [("Heading 1", 15), ("Heading 2", 12)]:
        style = doc.styles[name]
        style.font.name = "Microsoft YaHei"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor(46, 116, 181)


def para(doc, text, bold=False, color=None):
    p = doc.add_paragraph()
    r = p.add_run(text)
    set_run_font(r, bold=bold, color=color)


def bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        r = p.add_run(item)
        set_run_font(r)


def table(doc, headers, rows):
    t = doc.add_table(rows=1, cols=len(headers))
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    t.style = "Table Grid"
    for i, header in enumerate(headers):
        cell = t.rows[0].cells[i]
        cell.text = ""
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(header)
        set_run_font(r, size=8.5, bold=True)
        set_cell_shading(cell, "F2F4F7")
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    for row in rows:
        cells = t.add_row().cells
        for i, value in enumerate(row):
            cells[i].text = ""
            p = cells[i].paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT if len(str(value)) > 16 else WD_ALIGN_PARAGRAPH.CENTER
            r = p.add_run(str(value))
            set_run_font(r, size=8)
            cells[i].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    doc.add_paragraph()


def read_case_matrix(path):
    if not path.exists():
        return []
    with path.open(newline="", encoding="utf-8") as f:
        return list(csv.DictReader(f))


def main():
    parser = argparse.ArgumentParser()
    parser.add_argument("--out", required=True, type=Path)
    parser.add_argument("--case-matrix", type=Path, default=Path(__file__).resolve().parents[1] / "case_matrix.csv")
    args = parser.parse_args()
    args.out.parent.mkdir(parents=True, exist_ok=True)

    doc = Document()
    setup(doc)
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = title.add_run("KDebug XiangShan Benchmark v2 设计说明")
    set_run_font(r, size=18, bold=True, color=(11, 37, 69))
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = subtitle.add_run("三模型 repair-loop 对比：gpt-5.5 / glm-4.7 / qwen3.6-35b")
    set_run_font(r, size=10, color=(85, 85, 85))

    doc.add_heading("一、设计目标", level=1)
    para(doc, "v2 benchmark 用于评估模型在真实 XiangShan RTL、验证环境和混合故障中的闭环修复能力。PASS 不是静态定位成功，而是在 1 小时内修改允许范围内的代码并让原 fail workload 重新通过。")
    para(doc, "本版本修复了旧版的主要偏差：with_kdebug 必须有真实工具证据；环境 bug 可以修改脚本/配置；混合 bug 必须同时修 RTL 和环境；报告按 bug 域和 benchmark 层拆分。", bold=True)

    doc.add_heading("二、分组规则", level=1)
    bullets(doc, [
        "with_kdebug：只允许使用当前 case 由独立 KDebug 采集生成、且通过 manifest 校验的 plan 声明响应。",
        "without_kdebug：只能使用当前 case 的日志、允许范围内的源文件和普通文本搜索，禁止使用 kdebug、FSDB、KDB、Verdi、NPI/Tcl 动态查询。",
        "两组都禁止读取 answer_key_private.json，禁止跨 case majority diff，禁止修改 judge 脚本和 pass marker。",
    ])

    doc.add_heading("三、PASS 口径", level=1)
    bullets(doc, [
        "rtl case：要求 RTL 有有效修改，build/run/judge 通过。",
        "env case：要求环境/脚本/配置有有效修改；如果通过是靠 RTL 修改，则判规则违规。",
        "mixed case：要求 RTL 和环境均有有效修改，build/run/judge 通过。",
        "任何 case 如果 with_kdebug 组缺少工具证据，标记为 TOOL_EVIDENCE_MISSING；manifest、输入、哈希、调用或响应校验失败标记为 TOOL_EVIDENCE_INVALID。两者都不计为普通模型失败或成功。",
    ])

    doc.add_heading("四、Case 矩阵", level=1)
    rows = read_case_matrix(args.case_matrix)
    table(
        doc,
        ["Case", "Benchmark 层", "级别", "子系统", "Bug 域", "故障类别", "修复范围"],
        [
            [
                r.get("case_id", ""),
                r.get("benchmark_layer", ""),
                r.get("level", ""),
                r.get("subsystem", ""),
                r.get("bug_domain", ""),
                r.get("bug_class", ""),
                r.get("repair_scope", ""),
            ]
            for r in rows
        ],
    )

    doc.add_heading("五、公开 RTL 设计资料映射", level=1)
    table(
        doc,
        ["Case", "推荐设计资料"],
        [[r.get("case_id", ""), r.get("recommended_design_refs", "")] for r in rows],
    )

    doc.add_heading("六、交付物", level=1)
    bullets(doc, [
        "results.csv：记录所有 trial 的耗时、token、修改文件、工具证据、失败类别和最终状态。",
        "summary.md：按模型、工具组、bug 域、benchmark 层拆分的中文汇总。",
        "docx_out/*.docx：每个模型一份中文 Word 报告，加三模型汇总报告。",
        "screenshots/*.png：每个 trial 的终端截图证据。",
    ])

    doc.save(args.out)


if __name__ == "__main__":
    main()
